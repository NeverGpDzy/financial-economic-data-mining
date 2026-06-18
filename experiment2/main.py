"""Executable pipeline for Experiment 2.

Run from repository root:

    python -m experiment2.main
"""

from __future__ import annotations

import argparse
import io
import json
import sqlite3
import sys
from pathlib import Path

import pandas as pd

if sys.stdout.encoding and sys.stdout.encoding.lower().replace("-", "") != "utf8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from . import config
from .analysis import (
    build_herd_index,
    build_quality_checks,
    build_summary_statistics,
    build_top_herd_weeks,
    load_weekly_sentiment,
)
from .plots import generate_all_plots


def _save_df(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path, index=False, encoding="utf-8-sig")


def _json_safe(obj):
    if hasattr(obj, "item"):
        return obj.item()
    if isinstance(obj, (pd.Timestamp,)):
        return obj.strftime("%Y-%m-%d")
    return str(obj)


def _repo_rel(path: str | Path) -> str:
    p = Path(path)
    try:
        return p.relative_to(config.ROOT).as_posix()
    except ValueError:
        return p.as_posix()


def save_sqlite(
    weekly_sentiment: pd.DataFrame,
    herd: pd.DataFrame,
    quality: pd.DataFrame,
    stats: pd.DataFrame,
    top_weeks: pd.DataFrame,
) -> None:
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(config.DB_FILE) as conn:
        weekly_sentiment.to_sql("weekly_sentiment_input", conn, if_exists="replace", index=False)
        herd.to_sql("weekly_herd_index", conn, if_exists="replace", index=False)
        quality.to_sql("quality_checks", conn, if_exists="replace", index=False)
        stats.to_sql("summary_statistics", conn, if_exists="replace", index=False)
        top_weeks.to_sql("top_herd_weeks", conn, if_exists="replace", index=False)


def _markdown_table(df: pd.DataFrame, columns: list[str], max_rows: int | None = None) -> list[str]:
    view = df[columns].copy()
    if max_rows is not None:
        view = view.head(max_rows)
    headers = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join(["---"] * len(columns)) + " |"
    rows = [headers, sep]
    for _, row in view.iterrows():
        vals = []
        for col in columns:
            val = row[col]
            if isinstance(val, pd.Timestamp):
                vals.append(val.strftime("%Y-%m-%d"))
            elif isinstance(val, float):
                vals.append(f"{val:.6f}")
            else:
                vals.append(str(val))
        rows.append("| " + " | ".join(vals) + " |")
    return rows


def write_report(
    audit: dict,
    herd: pd.DataFrame,
    quality: pd.DataFrame,
    stats: pd.DataFrame,
    top_weeks: pd.DataFrame,
    plots: dict[str, str],
) -> None:
    config.REPORT_DIR.mkdir(parents=True, exist_ok=True)
    top1 = top_weeks.iloc[0]
    passed = int(quality["result"].sum())
    total_checks = int(len(quality))

    lines = [
        "# 实验二：金融非结构化数据情感分析报告",
        "",
        f"学生：{config.STUDENT_NAME}  ",
        f"学号：{config.STUDENT_ID}",
        "",
        "## 1. 作业要求理解",
        "",
        "实验二承接实验一输出的周度情绪指标，目标是构建可用于后续回归建模的周度羊群效应指数。核心逻辑是：先用正面情绪占比表示市场乐观程度，再用过去4周均值作为正常情绪基准，最后把“情绪是否异常”和“观点是否一边倒”合成为 H3 羊群效应指数。",
        "",
        "本实现采用以下口径：",
        "",
        "- `P_t = WeekPositive_t / (WeekPositive_t + WeekNegative_t)`，只在有明确正负方向的新闻中计算乐观占比。",
        "- `E(P_t)` 使用过去4周 `P_t` 的滚动均值，并且不使用当周数据，避免同周信息污染基准。",
        "- `H1t = P_t - E(P_t)`，衡量本周情绪相对历史基准的偏离。",
        "- `H2t = 1 - |WeekPositive_t - WeekNegative_t| / (WeekPositive_t + WeekNegative_t)`，作为多空分歧度；数值越小，说明正负观点越一边倒。",
        "- `H3t = Norm(|H1t|) * (1 - Norm(H2t))`，同时强调情绪反常幅度和一边倒强度。",
        "- `H2t_formula_raw` 和 `H3t_formula_raw` 保留任务书公式原文转写口径，用于审查对照。",
        "",
        "说明：指导书正文写明“H2 越小 → 大家一边倒 → 羊群效应越强”。为保持公式与解释一致，最终提交指标 `H2t` 采用分歧度，并在 H3 合成时使用 `1 - Norm(H2t)` 表示一致性强度；同时保留公式原文列，便于按任务书字面复核。",
        "",
        "## 2. 数据与预处理",
        "",
        f"- 输入文件：`{_repo_rel(config.INPUT_WEEKLY_SENTIMENT)}`。",
        f"- 输入周度样本：{audit['input_rows']} 行。",
        f"- 无正负情绪分母剔除：{audit['dropped_zero_denominator']} 行。",
        f"- 无历史基准剔除：{audit['dropped_no_history_baseline']} 行。",
        f"- 3σ异常值剔除：{audit['outlier_removed']} 行。",
        f"- 最终输出：{audit['output_rows']} 行，日期范围 {audit['date_start']} 至 {audit['date_end']}。",
        "",
        "## 3. 指标质量检查",
        "",
        f"质量检查通过 {passed}/{total_checks} 项。",
        "",
    ]
    lines += _markdown_table(quality, ["check", "result", "detail"])
    lines += [
        "",
        "主要指标描述统计如下：",
        "",
    ]
    lines += _markdown_table(stats, ["indicator", "count", "mean", "std", "min", "median", "max"])
    lines += [
        "",
        "## 4. 羊群效应结果",
        "",
        f"H3 最大值出现在 {pd.to_datetime(top1['trade_date']).strftime('%Y-%m-%d')}，H3={top1['H3t']:.6f}，当周 P_t={top1['P_t']:.6f}，H1={top1['H1t']:.6f}，H2分歧度={top1['H2t']:.6f}。",
        "",
        "H3 排名前10的交易周：",
        "",
    ]
    lines += _markdown_table(
        top_weeks,
        [
            "rank",
            "trade_date",
            "P_t",
            "E_P_t",
            "H1t",
            "H2t",
            "H2t_formula_raw",
            "H2t_strength",
            "H3t_formula_raw",
            "H3t",
            "WeekPositive",
            "WeekNegative",
            "NewsCount",
        ],
    )
    lines += [
        "",
        "## 5. 输出文件",
        "",
        "- `experiment2/assignment.md`：从最新版实验指导书中提取的实验二要求。",
        "- `data/experiment2/original/`：老师发布的最新版指导书与配套数据 zip。",
        "- `data/experiment2/raw/`：解压后的新闻数据和沪深300价格数据。",
        "- `outputs/experiment2/weekly_sentiment_input.csv`：实验一周度情绪输入副本。",
        "- `outputs/experiment2/weekly_herd_index.csv`：实验二核心输出，包含 `trade_date`、`H1t`、`H2t`、`H3t`。",
        "- `outputs/experiment2/experiment2.db`：SQLite 数据库，包含输入、指标、质量检查、描述统计和高羊群周表。",
        "- `outputs/experiment2/202331060205_丁致宇_实验二_羊群效应指数构建报告.docx`：可提交的 Word 版标准实验报告。",
        "- `outputs/experiment2/*.png`：指标时序、情绪对比、Top周和分布图。",
        "- `outputs/experiment2/AI交互记录.md`、`outputs/experiment2/AI代码审查与修复表.md`：AI辅助材料。",
        "- `outputs/experiment2/实验二代码附录.md`：本次实验二代码附录。",
        "",
        "## 6. 图表索引",
        "",
        f"- 羊群效应指标时序：`{_repo_rel(plots['herd_timeseries'])}`",
        f"- 情绪与羊群效应对比：`{_repo_rel(plots['sentiment_vs_herd'])}`",
        f"- H3最高交易周：`{_repo_rel(plots['top_herd_weeks'])}`",
        f"- 指标分布：`{_repo_rel(plots['indicator_distribution'])}`",
        f"- 指标表截图：`{_repo_rel(plots['herd_table_snapshot'])}`",
        f"- 核心程序截图：`{_repo_rel(plots['core_code_snippet'])}`",
        "",
        "## 7. 金融逻辑说明",
        "",
        "单纯的乐观或悲观并不等同于羊群效应。只有当本周情绪显著偏离过去基准，同时市场观点又明显向同一方向集中时，才更符合羊群行为的特征。H1 捕捉“反常”，H2 捕捉“是否分歧”，H3 则把两者合成，过滤掉普通情绪波动和多空分歧较大的噪声周。",
    ]

    text = "\n".join(lines)
    (config.OUTPUT_DIR / "experiment2_report.md").write_text(text, encoding="utf-8")
    (config.REPORT_DIR / "experiment2_report.md").write_text(text, encoding="utf-8")


def _add_docx_table(document, df: pd.DataFrame, columns: list[str], max_rows: int | None = None) -> None:
    view = df[columns].head(max_rows).copy() if max_rows else df[columns].copy()
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, col in enumerate(columns):
        table.rows[0].cells[idx].text = col
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            val = row[col]
            if isinstance(val, pd.Timestamp):
                text = val.strftime("%Y-%m-%d")
            elif isinstance(val, float):
                text = f"{val:.6f}"
            else:
                text = str(val)
            cells[idx].text = text


def write_docx_report(
    audit: dict,
    quality: pd.DataFrame,
    stats: pd.DataFrame,
    top_weeks: pd.DataFrame,
    plots: dict[str, str],
) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt

    path = config.OUTPUT_DIR / f"{config.STUDENT_ID}_{config.STUDENT_NAME}_实验二_羊群效应指数构建报告.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    document.add_heading("实验二：金融非结构化数据情感分析", level=0)
    document.add_paragraph(f"学生：{config.STUDENT_NAME}    学号：{config.STUDENT_ID}")
    document.add_paragraph("主题：基于市场情绪的金融羊群效应指数构建")

    document.add_heading("一、实验目标与口径", level=1)
    for text in [
        "读取实验一输出的周度情绪指标，构建 H1、H2、H3 周度羊群效应指数。",
        "P_t 按 WeekPositive / (WeekPositive + WeekNegative) 计算。",
        "E(P_t) 使用过去4周 P_t 滚动均值，不包含当周。",
        "H1t 衡量情绪偏离；H2t 为解释一致版分歧度，越小代表越一边倒；H3t 综合情绪异常强度与一致性强度。",
        "H2t_formula_raw、H3t_formula_raw 保留任务书公式原文转写口径，供审查对照。",
    ]:
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("二、数据处理摘要", level=1)
    for text in [
        f"输入周度样本：{audit['input_rows']} 行。",
        f"无历史基准剔除：{audit['dropped_no_history_baseline']} 行。",
        f"3σ异常值剔除：{audit['outlier_removed']} 行。",
        f"最终输出：{audit['output_rows']} 行，日期范围 {audit['date_start']} 至 {audit['date_end']}。",
    ]:
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("三、质量检查", level=1)
    _add_docx_table(document, quality, ["check", "result", "detail"])

    document.add_heading("四、描述统计", level=1)
    _add_docx_table(document, stats, ["indicator", "count", "mean", "std", "min", "median", "max"])

    document.add_heading("五、羊群效应结果", level=1)
    _add_docx_table(
        document,
        top_weeks,
        ["rank", "trade_date", "P_t", "E_P_t", "H1t", "H2t", "H2t_formula_raw", "H3t"],
        max_rows=10,
    )

    document.add_heading("六、运行结果图与截图", level=1)
    for key, caption in [
        ("herd_table_snapshot", "羊群效应指标时序表截图"),
        ("core_code_snippet", "系统主要程序截图"),
        ("herd_timeseries", "H1、H2一致性强度和H3时序图"),
        ("sentiment_vs_herd", "情绪占比与羊群效应对比图"),
    ]:
        document.add_paragraph(caption)
        document.add_picture(plots[key], width=Inches(6.5))

    document.add_heading("七、金融逻辑说明", level=1)
    document.add_paragraph(
        "单纯的乐观或悲观不等同于羊群效应。H1 捕捉本周情绪相对历史基准的异常程度，"
        "H2 捕捉正负观点是否明显一边倒，H3 将两者合成后，可以过滤普通情绪波动和多空分歧较大的噪声周。"
    )
    document.add_paragraph("完整代码见 outputs/experiment2/实验二代码附录.md。")
    document.save(path)
    return path


def write_ai_files(output_dir: Path) -> None:
    interaction = """# 实验二 AI交互记录

1. 资料整理：读取老师最新版实验指导书，定位“实验二 金融非结构化数据情感分析”章节，并提取为 Markdown。
2. 数据归档：复制老师发布的 Word、zip 原件，将配套原始数据放入 `data/experiment2/raw/`。
3. 要求分析：确认实验二依赖实验一输出的周度情绪表，核心指标为 P_t、E(P_t)、H1、H2、H3。
4. 公式实现：按正面/(正面+负面)重算 P_t；用过去4周 P_t 均值作为基准；计算 H1 情绪偏离。
5. H2方向处理：根据指导书“H2越小代表越一边倒”的解释，将 H2 实现为正负观点分歧度，并在 H3 中取反向强度。
6. 结果输出：生成 `weekly_herd_index.csv`、SQLite数据库、质量检查、描述统计、Top羊群周、指标表截图和核心程序截图。
7. 报告生成：生成 Markdown 报告、Word 版标准实验报告、AI代码审查表和代码附录，保证可提交材料完整。
"""
    audit = """# 实验二 AI代码审查与修复表

| 编号 | 审查发现 | 影响 | 修复动作 | 验证 |
| --- | --- | --- | --- | --- |
| 1 | 实验二需要独立可交付，不能继续混放在实验一输出中。 | 提交材料边界不清晰。 | 新建 `experiment2/`、`data/experiment2/`、`outputs/experiment2/`。 | `python -m experiment2.main` 可独立运行。 |
| 2 | H2文字解释为“越小越一边倒”，若直接用正负差额会与正向羊群强度冲突。 | H3方向可能反了。 | 将 H2 定义为分歧度 `1 - abs(pos-neg)/(pos+neg)`，H3 使用 `1 - Norm(H2)`。 | `quality_checks.csv` 验证 H2、H3范围均为0~1。 |
| 3 | 4周基准若包含当周 P_t 会产生同周信息污染。 | H1偏离度被低估。 | 先 `shift(1)` 再做4周 rolling mean。 | 第一周因无历史基准被剔除，报告记录剔除数量。 |
| 4 | 输出只有CSV不满足报告要求。 | 缺少图表、数据库和AI审查材料。 | 同步输出 SQLite、PNG图表、Markdown报告、Word报告、AI交互记录、代码审查表和代码附录。 | `outputs/experiment2/` 生成完整文件。 |
"""
    (output_dir / "AI交互记录.md").write_text(interaction, encoding="utf-8")
    (output_dir / "AI代码审查与修复表.md").write_text(audit, encoding="utf-8")


def write_code_appendix(output_dir: Path) -> None:
    files = [
        config.PACKAGE_DIR / "config.py",
        config.PACKAGE_DIR / "analysis.py",
        config.PACKAGE_DIR / "plots.py",
        config.PACKAGE_DIR / "main.py",
    ]
    lines = ["# 实验二代码附录", ""]
    for path in files:
        rel = path.relative_to(config.ROOT).as_posix()
        lines += [f"## {rel}", "", "```python", path.read_text(encoding="utf-8").rstrip(), "```", ""]
    (output_dir / "实验二代码附录.md").write_text("\n".join(lines), encoding="utf-8")


def write_data_description(audit: dict) -> None:
    config.DATA_DIR.mkdir(parents=True, exist_ok=True)
    text = f"""# 实验二配套数据说明

## 原始资料

- `original/金融数据挖掘实验指导书 20260617.docx`：老师发布的最新版实验指导书原件。
- `original/实验配套数据.zip`：老师发布的实验配套数据压缩包。
- `raw/新闻数据.xls`：金融新闻文本数据，与实验一原始数据一致。
- `raw/沪深300日价格指数.xls`：沪深300日价格数据，与实验一原始数据一致。

## 实验二输入

- 直接输入：`outputs/experiment1/weekly_sentiment.csv`。
- 输入字段：`week`、`WeekPositive`、`WeekNeutral`、`WeekNegative`、`NewsCount`、`P_t`。
- 本实验重新按 `WeekPositive / (WeekPositive + WeekNegative)` 计算 `P_t`，保证与实验二公式一致。
- 输出同时保留 `H2t_formula_raw`、`H3t_formula_raw`，用于对照任务书公式原文转写口径。

## 处理摘要

- 输入周度样本：{audit.get('input_rows')} 行。
- 无正负情绪分母剔除：{audit.get('dropped_zero_denominator')} 行。
- 无历史基准剔除：{audit.get('dropped_no_history_baseline')} 行。
- 3σ异常值剔除：{audit.get('outlier_removed')} 行。
- 最终输出周度羊群指标：{audit.get('output_rows')} 行。
- 输出日期范围：{audit.get('date_start')} 至 {audit.get('date_end')}。

## 主要输出

- `outputs/experiment2/weekly_herd_index.csv`：实验二羊群效应指标表。
- `outputs/experiment2/experiment2.db`：SQLite数据库。
- `outputs/experiment2/experiment2_report.md`：实验二报告。
- `outputs/experiment2/202331060205_丁致宇_实验二_羊群效应指数构建报告.docx`：Word版标准实验报告。
"""
    (config.DATA_DIR / "数据说明.md").write_text(text, encoding="utf-8")


def run() -> dict:
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    config.REPORT_DIR.mkdir(parents=True, exist_ok=True)

    print("实验二：读取实验一周度情绪指标...")
    weekly_sentiment = load_weekly_sentiment()

    print("构建 H1、H2、H3 羊群效应指标...")
    herd, audit = build_herd_index(weekly_sentiment)
    quality = build_quality_checks(herd)
    stats = build_summary_statistics(herd)
    top_weeks = build_top_herd_weeks(herd)

    print("保存 CSV 和 SQLite 数据库...")
    _save_df(weekly_sentiment, config.OUTPUT_DIR / "weekly_sentiment_input.csv")
    _save_df(herd, config.OUTPUT_DIR / "weekly_herd_index.csv")
    _save_df(quality, config.OUTPUT_DIR / "quality_checks.csv")
    _save_df(stats, config.OUTPUT_DIR / "summary_statistics.csv")
    _save_df(top_weeks, config.OUTPUT_DIR / "top_herd_weeks.csv")
    save_sqlite(weekly_sentiment, herd, quality, stats, top_weeks)

    print("生成图表、报告和AI辅助材料...")
    plots = generate_all_plots(herd, top_weeks, config.OUTPUT_DIR)
    write_ai_files(config.OUTPUT_DIR)
    write_code_appendix(config.OUTPUT_DIR)
    write_report(audit, herd, quality, stats, top_weeks, plots)
    docx_report = write_docx_report(audit, quality, stats, top_weeks, plots)
    write_data_description(audit)

    summary = {
        "student": {"name": config.STUDENT_NAME, "id": config.STUDENT_ID},
        "audit": audit,
        "quality_passed": int(quality["result"].sum()),
        "quality_total": int(len(quality)),
        "top_herd_week": top_weeks.iloc[0].to_dict(),
        "plots": plots,
        "docx_report": str(docx_report),
    }
    (config.OUTPUT_DIR / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2, default=_json_safe),
        encoding="utf-8",
    )

    print(f"完成。输出目录：{config.OUTPUT_DIR}")
    return summary


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="实验二：基于市场情绪的金融羊群效应指数构建")
    return parser.parse_args()


def main() -> None:
    parse_args()
    run()


if __name__ == "__main__":
    main()
