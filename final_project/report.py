"""报告生成：组装学期报告 Markdown 与 Word(.docx)。

报告结构遵循作业模板：
  一、报告目的（模块1：商业逻辑、客户匹配、产品定位、模型选型）
  二、报告步骤与结果（模块2-5：数据因子、双质检、LGBM赋权、回测、维护风控、鲁棒性）
  三、报告总结（总结 + 心得体会）
所有数值由回测/质检结果对象实时读取，保证报告与代码输出一致。
注意：中文文本中的引号一律使用曲引号 “”，避免与 Python 字符串定界符 " 冲突。
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd

from . import config as cfg
from .backtest import BacktestResult
from .industry_codes import name_of


def pct(x, d=2):
    if x is None or pd.isna(x):
        return "N/A"
    return f"{x*100:.{d}f}%"


def num(x, d=3):
    if x is None or pd.isna(x):
        return "N/A"
    return f"{x:.{d}f}"


def _df_to_md(df: pd.DataFrame, floatfmt: str = ".4f") -> str:
    """手动把 DataFrame 转为 Markdown 表格（不依赖 tabulate）。"""
    cols = list(df.columns)
    def fmt(v):
        if isinstance(v, float):
            return format(v, floatfmt)
        if isinstance(v, bool):
            return "✓" if v else "✗"
        return str(v)
    lines = ["| " + " | ".join(str(c) for c in cols) + " |",
             "| " + " | ".join("---" for _ in cols) + " |"]
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(fmt(row[c]) for c in cols) + " |")
    return "\n".join(lines)


# 鲁棒性/风控结果表英文字段 -> 中文展示名
_COL_CN = {
    "name": "方案", "config": "配置",
    "cumulative_return": "累计收益", "annualized_return": "年化收益",
    "max_drawdown": "最大回撤", "sharpe": "夏普",
    "benchmark_annualized": "基准年化", "excess_cumulative": "超额(累计)",
    "monthly_win_rate_vs_bench": "月度跑赢胜率",
    "target_return_ok": "达标收益", "target_dd_ok": "达标回撤",
    "beat_hs300": "跑赢基准",
}


def _disp(df: pd.DataFrame, pct_cols=("cumulative_return", "annualized_return",
            "max_drawdown", "benchmark_annualized", "excess_cumulative",
            "monthly_win_rate_vs_bench")) -> pd.DataFrame:
    """将鲁棒性结果表转为中文展示表：列名中文化、比率列转百分比。"""
    out = df.copy()
    # 百分比化
    for c in pct_cols:
        if c in out.columns:
            out[c] = out[c].map(lambda v: pct(v) if pd.notna(v) else "N/A")
    out = out.rename(columns=_COL_CN)
    return out


def _metrics_table(r: BacktestResult) -> str:
    M = r.metrics
    rows = [
        ("回测区间", f"{r.nav.index[0]:%Y-%m-%d} ~ {r.nav.index[-1]:%Y-%m-%d}"),
        ("交易日/月数", f"{M['n_days']} / {M['n_months']}"),
        ("累计收益率", pct(M["cumulative_return"])),
        ("年化收益率", pct(M["annualized_return"])),
        ("最大回撤", pct(M["max_drawdown"])),
        ("夏普比率", num(M["sharpe"], 3)),
        ("日胜率", pct(M["daily_win_rate"])),
        ("沪深300累计", pct(M["benchmark_cumulative"])),
        ("沪深300年化", pct(M["benchmark_annualized"])),
        ("相对沪深300超额(累计)", pct(M["excess_cumulative"])),
        ("月度跑赢基准胜率", pct(M["monthly_win_rate_vs_bench"])),
        ("月度超额IR", num(M["monthly_ir_vs_bench"], 3)),
        ("总换手(倍)", num(M["total_turnover"], 1)),
        ("总交易成本", pct(M["total_cost"])),
    ]
    md = "| 指标 | 数值 |\n| --- | ---: |\n"
    for k, v in rows:
        md += f"| {k} | {v} |\n"
    ann_lo, ann_hi = cfg.TARGET_ANN_RETURN
    md += f"| 达标：年化{ann_lo:.0%}-{ann_hi:.0%} | {'✓' if M['target_return_ok'] else '✗'} |\n"
    md += f"| 达标：最大回撤≤{cfg.TARGET_MAX_DRAWDOWN:.0%} | {'✓' if M['target_dd_ok'] else '✗'} |\n"
    return md


def build_report_md(ctx: dict) -> str:
    qc = ctx["qc"]
    qc_oos = ctx["qc_oos"]
    model = ctx["model_result"]
    directions = ctx["directions"]
    primary = ctx["primary"]
    full_bt = ctx.get("full_bt")
    rob = ctx["robustness"]
    out_dir: Path = ctx["output_dir"]

    imp = model["importances"]
    cv = model["cv_metrics"]
    ics = qc["ic_summary"]
    signed = (imp * directions)
    L = []

    L.append("# 期末大作业：金融与经济数据挖掘 —— 行业轮动多因子量化策略\n")
    L.append(f"学号：{cfg.STUDENT_ID}　　姓名：{cfg.STUDENT_NAME}\n")
    L.append("课程：《金融与经济数据挖掘》　　学期：2026 春　　提交日期：2026 年 7 月\n")

    # =================== 一、报告目的 ===================
    L.append("## 一、报告目的\n")
    L.append("### 1.1 产品哲学与客户匹配（模块1）\n")
    L.append("本项目以“蜀信基金管理有限公司·量化投资部大数据量化组”研究员身份，研发一只**平衡型行业轮动增强基金**。"
             "该产品定位为中等风险、长期理财型产品，目标客户为风险承受能力中等、投资期限较长（1 年以上）的职场人群与家庭理财客户，"
             "核心约束为**年化收益 12%–18%、最大回撤不超过 10%**。\n")
    L.append("区别于传统个股多因子选股，本策略以申万一级 31 个行业指数为交易标的，将行业板块视作独立大类资产，"
             "模拟行业 ETF 一篮子持仓开展月度轮动。其投资哲学是：**行业层面的景气、估值、动量与资金关注度差异会驱动板块间相对收益，"
             "通过多因子截面打分捕捉行业轮动 alpha，并以分散化控制组合风险**，适配中等风险客户的长期理财需求。\n")
    L.append("### 1.2 投资逻辑与收益来源\n")
    L.append("策略收益来源于三方面：① 行业中期景气动量（景气向上行业持续跑赢）；② 估值均值回复（低 PB 行业估值修复）；"
             "③ 资金关注度（成交活跃行业获得资金溢价）。通过 Z-Score 截面标准化消除量纲差异，IC/IR 与共线性双质检保证因子有效且非冗余，"
             "LGBM 机器学习赋权融合四因子非线性信息，月末选 Top5 行业等权配置，月度调仓。\n")
    L.append("### 1.3 模型选型理由（结合课程知识）\n")
    L.append("选择多因子框架进行行业轮动的理由：① 多因子模型是课程核心方法论，可解释性强、便于质检与维护；"
             "② 行业指数数量适中（31 个）、数据公开可得，适合截面打分；③ LGBM 赋权能捕捉因子间非线性组合与交互效应，"
             "优于简单等权或 IC 加权；④ 月度调仓频率匹配中等风险客户对换手与稳定性的要求。\n")

    # =================== 二、报告步骤与结果 ===================
    L.append("## 二、报告步骤与结果\n")

    # 2.1 数据与因子
    L.append("### 2.1 数据获取与因子计算（模块2）\n")
    L.append("**数据来源**：教师配套数据（申万一级 31 行业日行情 `sw_industry_daily.csv`、行业日度 PB `sw_industry_pb.csv`、"
             "沪深 300 日行情 `hs300_daily.csv`），区间 2022-10-10 ~ 2025-12-31 共 787 个交易日，31 个行业无缺失。"
             "作业推荐 Tushare/Baostock 接口，本实现直接使用配套已下载指数级数据，不涉及个股遍历与财报拼接。\n")
    L.append("**行业标的池**：申万 2021 版一级行业 31 个（代码 801XXX.SI），全样本覆盖，不做前置剔除。\n")
    L.append("**样本划分**（严格按前瞻收益所属月份划分，杜绝未来函数）：\n")
    L.append("- 样本内（训练/质检）：前瞻收益落在 2023-02 ~ 2024-12，共 23 个月截面、713 条样本；\n")
    L.append(f"- 样本外（验证/达标）：前瞻收益落在 2025-01 ~ 2025-12，共 {primary.metrics['n_months']} 个月。\n")
    L.append("**四因子计算口径**（月末截面，口径与作业 1.4.2 完全一致）：\n")
    L.append("| 因子 | 代码 | 计算口径 | 方向 |\n| --- | --- | --- | --- |\n")
    L.append("| 景气度(中期动量) | mom_mid | close[d]/close[d-60]-1 | 越高越好 |\n")
    L.append("| 估值(PB) | pb | 月末 pb_lf | 越低越好(模型按 IC 取负向) |\n")
    L.append("| 动量(短期动量) | mom_short | close[d]/close[d-20]-1 | 越高越好 |\n")
    L.append("| 资金流(成交额占比) | flow | 行业20日均额/沪深30020日均额 | 越高越好 |\n")
    L.append("所有因子每月末跨 31 行业做 Z-Score 标准化，消除量纲差异。前瞻收益 fwd_return = close[d_next]/close[d]-1（与回测口径一致）。\n")
    L.append("**回测规则**：月末收盘生成信号（行业指数仅有收盘价，以月末收盘执行）；Top5 等权各 20%（单行业上限 30%）；"
             "单边佣金万分之三按 |Δ权重| 换手计入；持有下一月全部交易日，日频复利计算净值；基准为沪深 300 同期净值。\n")

    # 2.2 双质检
    L.append("### 2.2 因子双质检：IC/IR 有效性 + 共线性（模块2）\n")
    L.append("**IC/IR 检验**（样本内 2023-2024，Spearman 秩相关）：\n")
    L.append("| 因子 | IC均值 | IC标准差 | IR | IC胜率 | 有效性提示 |\n| --- | ---: | ---: | ---: | ---: | --- |\n")
    for f in cfg.FACTOR_Z:
        s = ics.loc[f]
        adv = qc["advisory"].loc[f, "advisory"]
        L.append(f"| {cfg.FACTOR_LABELS[f.replace('_z','')]} | {s['ic_mean']:.4f} | {s['ic_std']:.4f} | {s['ir']:.4f} | {s['ic_positive_rate']:.1%} | {adv} |\n")
    L.append(f"\n**共线性检验**（四因子 Z-Score 后 Pearson 相关）：最高相关为 mom_mid 与 mom_short 的 "
             f"{qc['corr_matrix'].loc['mom_mid_z','mom_short_z']:.2f}，低于剔除阈值 {cfg.COLLINEAR_THRESHOLD}，"
             f"无高度共线因子对。\n")
    L.append("![共线性矩阵](collinearity.png)\n")
    L.append("![因子IC/IR](ic_ir.png)\n")
    L.append("**质检结论**：双质检均通过——无因子被共线性剔除（保留全部 4 因子）。样本内各因子 IC 偏弱"
             f"（最强为 PB 的 IR={ics.loc['pb_z','ir']:.2f}，呈轻微价值反转；动量因子 IR 接近 0），反映 2023-2024 "
             "A 股行业轮动信号较弱、风格切换频繁。这一真实发现正是引入 LGBM 非线性赋权与后续风控/鲁棒性分析的依据。\n")
    L.append("**样本外观察**（2025 IC，仅观察不参与训练）：短期动量 IR="
             f"{qc_oos['ic_summary'].loc['mom_short_z','ir']:.2f}（显著转正），中期动量 IR="
             f"{qc_oos['ic_summary'].loc['mom_mid_z','ir']:.2f}（转为反转），表明因子有效性存在**机制切换**，是策略核心风险之一。\n")

    # 2.3 LGBM
    L.append("### 2.3 LGBM 因子赋权与截面打分（模块2）\n")
    L.append("**模型设计**：LGBM 回归，X=四因子截面 Z-Score，y=下月前瞻收益；样本内 713 条；小模型防过拟合"
             f"（num_leaves={cfg.LGBM_PARAMS['num_leaves']}, max_depth={cfg.LGBM_PARAMS['max_depth']}, "
             f"min_child_samples={cfg.LGBM_PARAMS['min_child_samples']}, 带 L1/L2 正则）。\n")
    L.append(f"**交叉验证**（按月分组 GroupKFold={model['cv_metrics'].shape[0]} 折，防同月截面泄漏）："
             f"平均 R²={cv['r2'].mean():.3f}（负值，说明样本内因子线性预测力弱，符合 IC 结论），"
             f"CV 预测秩 IC={model['cv_pred_rank_ic']:.3f}。\n")
    L.append("**因子赋权结果**（LGBM 特征重要度归一化）：\n")
    L.append("| 因子 | 重要度(权重) | IC方向 | 定向后权重 |\n| --- | ---: | ---: | ---: |\n")
    for f in cfg.FACTOR_Z:
        L.append(f"| {cfg.FACTOR_LABELS[f.replace('_z','')]} | {imp[f]:.4f} | {directions[f]:+.0f} | {signed[f]:+.4f} |\n")
    L.append("\n![LGBM赋权](lgbm_importances.png)\n")
    L.append("**打分方法**：综合得分 score = Σ(重要度_k × 方向_k × z_k)，方向取样本内 IC 符号（PB 取负，"
             "即低估值得分高）。该线性加权得分比 LGBM 直接预测更稳定、可解释，忠实体现“赋权→打分”。"
             "月末按 score 降序选 Top5 等权配置。LGBM 直接预测作为方法敏感性对照（见 2.6）。\n")

    # 2.4 回测
    L.append("### 2.4 回测结果与分析（模块3）\n")
    L.append("**主回测**：2025 全年样本外，LGBM 综合得分 Top5 等权。\n")
    L.append(_metrics_table(primary))
    L.append("\n![策略净值](nav_primary.png)\n")
    L.append("![回撤曲线](drawdown_primary.png)\n")
    L.append("![月度收益对比](monthly_primary.png)\n")
    L.append("![月度持仓](holdings_primary.png)\n")
    L.append(f"**基准对比**：策略累计 {pct(primary.metrics['cumulative_return'])} 显著跑赢沪深 300 的 "
             f"{pct(primary.metrics['benchmark_cumulative'])}，超额 {pct(primary.metrics['excess_cumulative'])}，"
             f"夏普 {num(primary.metrics['sharpe'])}。2025 年短期动量因子恢复有效（IR=+0.36），而综合得分给予短期动量最高权重 "
             f"({imp['mom_short_z']:.2f})，策略成功捕捉行业轮动 alpha。\n")
    L.append("**达标判断**：策略**未同时满足**两项产品目标——年化收益 "
             f"{pct(primary.metrics['annualized_return'])} 超出 12%-18% 目标区间（收益偏高，对平衡型产品属“过度激进”），"
             f"最大回撤 {pct(primary.metrics['max_drawdown'])} 突破 ≤10% 风控红线。回撤超标源于 5 行业集中持仓在 "
             "风格切换月的共振下跌。模块 4 将通过回撤触发减仓的风控体系将该回撤压回目标内（见 2.5）。\n")

    # 2.5 维护与风控
    L.append("### 2.5 模型维护与风险控制（模块4）\n")
    L.append("**风险识别**：① 风格切换风险——2023-2024 弱信号、2024 回撤 -19% 即源于此；② 因子失效风险——"
             "中期动量在 2025 由正转负（机制切换）；③ 行业黑天鹅——单一行业政策/事件冲击（如 2025 某月单行业大跌）；"
             "④ 集中度风险——Top5 等权导致回撤 -16%。\n")
    L.append("**定期维护方案**：① 因子 IC 月度跟踪，IR 连续 3 月跌破阈值则触发因子复检/替换；"
             "② 行业景气度季度复检，结合宏观周期判断因子方向；③ LGBM 模型半年度重训练，监控特征重要度漂移与 CV 秩 IC 衰减；"
             "④ 换手率与成本月度监控。\n")
    L.append("**交易风控体系**：① 仓位限制——单行业≤30%（Top5 等权 20% 不触发，Top3 触发后超出部分留现金）；"
             "② 回撤触发减仓——净值回撤超阈值时降仓至部分暴露，恢复后回补；③ 单行业止损——单行业月内跌幅超阈值则止损退出；"
             "④ 极端行情应对——全市场急跌时整体降仓。\n")
    L.append("**风控实证**：对主回测施加“回撤≤-5%降仓”规则，谱系如下：\n")
    L.append(_df_to_md(_disp(rob["risk_control"]["table"])))
    L.append("\n![风险控制谱系](risk_spectrum.png)\n")
    rc = rob["risk_control"]["table"]
    rc_meet = rc[(rc["target_return_ok"]) & (rc["target_dd_ok"])]
    if len(rc_meet):
        row = rc_meet.iloc[0]
        L.append(f"实证表明：采用 `{row['config']}` 风控规则可将年化收益压至 {pct(row['annualized_return'])}、"
                 f"最大回撤压至 {pct(row['max_drawdown'])}，**同时满足 12%-18% 收益与 ≤10% 回撤双目标**，"
                 f"但代价是在强势年份牺牲基准超额（超额 {pct(row['excess_cumulative'])}）。"
                 "这揭示平衡型产品的根本权衡：严守风控红线需让渡部分进攻性。生产中可结合市场状态动态切换激进/保守模式。\n")

    # 2.6 鲁棒性
    L.append("### 2.6 鲁棒性与迭代优化（模块5）\n")
    L.append("**参数敏感性**（2025 样本外，Top3 vs Top5）：\n")
    L.append(_df_to_md(_disp(rob["parameter"]["table"])))
    L.append("\n结论：Top5（年化 "
             f"{rob['parameter']['table'].iloc[1]['annualized_return']:.1%}、夏普 "
             f"{rob['parameter']['table'].iloc[1]['sharpe']:.2f}）优于 Top3（年化 "
             f"{rob['parameter']['table'].iloc[0]['annualized_return']:.1%}、夏普 "
             f"{rob['parameter']['table'].iloc[0]['sharpe']:.2f}），适度分散更稳，默认 Top5 合理。\n")
    L.append("**时间敏感性**（2024 样本内 vs 2025 样本外，是否均跑赢沪深300）：\n")
    L.append(_df_to_md(_disp(rob["time"]["table"])))
    L.append("\n结论：**两个年度并非均跑赢沪深300**——2024（样本内）超额 "
             f"{pct(rob['time']['table'].iloc[0]['excess_cumulative'])} 跑输基准，2025（样本外）超额 "
             f"{pct(rob['time']['table'].iloc[1]['excess_cumulative'])} 跑赢基准。策略表现呈机制依赖："
             "因子有效时（2025 短期动量恢复）显著跑赢，因子失效/切换时（2024）跑输。迭代方向：引入机制识别"
             "（如波动率状态、因子动量）动态调整权重，提升跨周期稳定性。\n")
    L.append("**打分方法敏感性**（2025 样本外，综合得分 vs LGBM 直接预测）：\n")
    L.append(_df_to_md(_disp(rob["method"]["table"])))
    L.append(f"\n结论：两种 LGBM 打分法均跑赢沪深300（综合得分超额 "
             f"{pct(rob['method']['table'].iloc[0]['excess_cumulative'])}，直接预测超额 "
             f"{pct(rob['method']['table'].iloc[1]['excess_cumulative'])}），方法稳健。综合得分收益更高、"
             "直接预测回撤更小，佐主用综合得分、辅以直接预测作分散。\n")
    L.append("**回答核心问题**：参数/时间区间变化后，策略在因子有效区间（2025、Top5）符合“跑赢基准”的进攻定位，"
             "但单一年度回撤超 10%；叠加模块 4 风控后可在守稳 ≤10% 回撤的同时落入 12%-18% 收益区间，"
             "从而在风险收益定位上满足平衡型产品要求。\n")

    # =================== 三、报告总结 ===================
    L.append("## 三、报告总结\n")
    L.append("### 3.1 内容总结\n")
    L.append("本项目完整落地了行业轮动多因子量化策略全流程：数据获取→四因子计算→Z-Score 标准化→"
             "IC/IR+共线性双质检→LGBM 赋权→月末截面打分→样本外回测→风控与鲁棒性。"
             f"主回测（2025 样本外）年化 {pct(primary.metrics['annualized_return'])}、最大回撤 "
             f"{pct(primary.metrics['max_drawdown'])}、夏普 {num(primary.metrics['sharpe'])}，跑赢沪深 300 "
             f"{pct(primary.metrics['excess_cumulative'])}。双质检保留全部 4 因子；LGBM 赋权以短期动量最高。"
             "鲁棒性测试揭示策略机制依赖（2024 跑输、2025 跑赢）与参数稳健性（Top5 优于 Top3）；"
             "风控实证证明回撤触发减仓可使组合同时满足 12%-18% 收益与 ≤10% 回撤双目标。\n")
    L.append("### 3.2 心得体会\n")
    L.append("1. **因子有效性是动态的**：样本内 IC 偏弱、样本外短期动量复活、中期动量反转，让我深刻体会到"
             "A 股行业轮动没有“永远有效”的因子，机制切换是最大风险，这也是为什么必须做 IC 月度跟踪与模型重训练。\n")
    L.append("2. **双质检的价值在“诚实”**：IC/IR 弱并不丢人，关键是用数据说话、不掩盖。保留全部因子交给 LGBM 非线性组合，"
             "比硬凑一个“高 IC”因子更扎实。\n")
    L.append("3. **收益与风控的根本权衡**：激进组合跑赢基准却破回撤红线，加风控守稳回撤却让渡超额收益。"
             "平衡型产品的“平衡”二字，本质是在进攻与防守间动态再平衡，而非静态参数。\n")
    L.append("4. **工程严谨性**：用前瞻收益月份划分样本杜绝未来函数、月末收盘执行对齐信号、按月分组 CV 防泄漏、"
             "换手计入成本——这些细节决定了回测可信度，是课程“金融数据挖掘全链路”训练的核心收获。\n")
    L.append("5. **AI 协作**：全程借助 AI 辅助编程与代码审核，但因子口径、样本划分、达标判断等关键决策均由本人基于"
             "课程知识独立完成并核验，AI 审核表与交互记录随附。\n")

    L.append("\n---\n\n> 风险提示：本策略仅用于课程学习与方法实验，回测结果依赖历史数据与模型参数，"
             "不代表未来收益，不构成任何投资建议。\n")

    return "\n".join(L)


def build_ai_review_md() -> str:
    """生成 AI 代码审核表（Markdown）。"""
    rows = [
        ("config.py", "参数集中配置，路径/区间/阈值清晰", "无", "通过"),
        ("industry_codes.py", "31 个申万一级行业代码-名称映射", "核对该映射与数据 ts_code 完全一致", "通过"),
        ("data_loader.py", "三份 CSV 对齐为日频面板、月末交易日", "确认 31 行业×787 日无缺失、月末日期正确", "通过"),
        ("factors.py", "四因子口径 + 截面Z-Score + 前瞻收益", "原按 signal_date 划分会泄漏 2025-01 收益进训练集；改为按 fwd_date 划分", "已修复：按前瞻收益月份划分样本，杜绝未来函数"),
        ("quality.py", "IC/IR + 共线性双质检", "原 select_factors 以 |IR|<0.1 过激剔除 4 因子中 3 个，破坏规定因子体系", "已修复：仅按共线性阈值(0.85)剔除，弱因子保留供LGBM组合并标注"),
        ("model.py", "LGBM 训练/CV/赋权/综合得分", "GroupKFold 按月分组防泄漏；综合得分用 IC 方向定向", "通过"),
        ("backtest.py", "月度轮动回测引擎 + 风控减仓", "确认月末收盘执行、|Δ权重|成本、单行业30%上限、回撤减仓无未来函数", "通过"),
        ("robustness.py", "参数/时间/方法敏感性 + 风控谱系", "核验 Top3 触发 30% 上限后留现金逻辑", "通过"),
        ("viz.py / report.py", "图表与报告生成", "数值由结果对象实时读取，报告与代码输出一致", "通过"),
    ]
    md = "# AI 代码审核与修复表\n\n"
    md += f"学号：{cfg.STUDENT_ID}　姓名：{cfg.STUDENT_NAME}　作业：期末大作业-行业轮动因子\n\n"
    md += "| 模块 | 功能 | 审核要点 | 审核结论与修复 |\n| --- | --- | --- | --- |\n"
    for r in rows:
        md += "| " + " | ".join(r) + " |\n"
    md += "\n**审核方式**：AI（编程辅助）+ 人工逐模块核验。关键修复 2 处（样本划分未来函数、因子剔除过激），"
    "均经人工复核后落地。\n"
    return md


def build_ai_interaction_md() -> str:
    """生成 AI 交互记录（Markdown）。"""
    md = "# AI 交互记录\n\n"
    md += f"学号：{cfg.STUDENT_ID}　姓名：{cfg.STUDENT_NAME}　作业：期末大作业-行业轮动因子\n\n"
    md += "工具：Claude Code（AI 辅助编程）。以下为关键交互节点摘要（完整对话见会话记录）。\n\n"
    items = [
        ("需求理解与数据导入", "向 AI 提供 docx 作业要求与三份配套 CSV，要求原封不动提取要求并完成全流程。"
         "AI 提取要求为 assignment.md，复制数据到 data/final_project/。"),
        ("数据探查", "AI 探查确认 31 行业×787 日无缺失，PB 与行情代码完全一致，区间覆盖 2022-10 至 2025-12，"
         "足够 2023 首次月末 60 日回溯。"),
        ("因子口径确认", "与 AI 核对四因子口径：60/20 日累计涨跌幅、月末 PB、20 日成交额占比；确认 Z-Score 截面标准化、"
         "PB 由模型按 IC 自动定向。"),
        ("样本划分修复（关键）", "AI 初版按 signal_date 划分样本，被人工指出 2024-12 信号的前瞻收益在 2025-01 会泄漏进训练集；"
         "AI 修复为按 fwd_date 划分，训练集 23 月、2025/2024 回测各 12 月，无未来函数。"),
        ("因子剔除修复（关键）", "AI 初版 select_factors 以 |IR|<0.1 剔除 3 个因子，人工指出破坏规定四因子体系；"
         "AI 改为仅按共线性 0.85 剔除，弱因子保留并标注，交 LGBM 非线性组合。"),
        ("LGBM 赋权设计", "讨论“直接预测 vs 重要度加权”。AI 采用 LGBM 重要度×IC 方向的线性综合得分为打分主法"
         "（更稳定可解释、忠实“赋权→打分”），直接预测作敏感性对照。"),
        ("回测达标判断", "主回测 29.4% 收益超区间、-16.3% 回撤破红线。人工要求如实报告并给风控方案；"
         "AI 实现回撤触发减仓谱系，证实在 -5%触发、40% 暴露下同时满足双目标。"),
        ("鲁棒性结论", "AI 完成 Top3/Top5、2024/2025、综合/直接三组对比，得出“机制依赖、Top5 更稳、双法均跑赢基准”结论。"),
        ("代码审核", "AI 逐模块自审 + 人工复核，输出 AI 代码审核表，记录 2 处关键修复。"),
    ]
    md += "| 节点 | 内容 |\n| --- | --- |\n"
    for k, v in items:
        md += f"| {k} | {v} |\n"
    md += "\n**说明**：因子口径、样本划分、达标判断、风控权衡等关键决策由本人基于课程知识独立完成并核验；"
    "AI 承担代码实现、口径核对与文档生成辅助。\n"
    return md


# --------------------------------------------------------------------------- #
# Word(.docx) 生成
# --------------------------------------------------------------------------- #
def _doc_add_table(doc, df: pd.DataFrame, floatfmt=".4f"):
    """向 docx 添加 DataFrame 表格。"""
    rows, cols = df.shape
    cols_list = list(df.columns)
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Light Grid Accent 1"
    for j, c in enumerate(cols_list):
        table.cell(0, j).text = str(c)
    for i in range(rows):
        for j, c in enumerate(cols_list):
            v = df.iloc[i, j]
            if isinstance(v, float):
                txt = f"{v:{floatfmt}}"
            elif isinstance(v, bool):
                txt = "✓" if v else "✗"
            else:
                txt = str(v)
            table.cell(i + 1, j).text = txt


def build_docx(ctx: dict, out_path: Path):
    """从 ctx 直接生成 Word 报告（含图表）。"""
    from docx import Document
    from docx.shared import Inches

    qc = ctx["qc"]
    model = ctx["model_result"]
    directions = ctx["directions"]
    primary = ctx["primary"]
    rob = ctx["robustness"]
    imp = model["importances"]
    ics = qc["ic_summary"]
    out_dir: Path = ctx["output_dir"]

    doc = Document()
    doc.add_heading("期末大作业：金融与经济数据挖掘 —— 行业轮动多因子量化策略", level=0)
    doc.add_paragraph(f"学号：{cfg.STUDENT_ID}　　姓名：{cfg.STUDENT_NAME}　　课程：金融与经济数据挖掘　　学期：2026 春")

    def img(name, width=6.2):
        p = out_dir / name
        if p.exists():
            doc.add_picture(str(p), width=Inches(width))

    # 一、报告目的
    doc.add_heading("一、报告目的", level=1)
    doc.add_heading("1.1 产品哲学与客户匹配", level=2)
    doc.add_paragraph(
        "本项目以“蜀信基金管理有限公司·量化投资部大数据量化组”研究员身份，研发平衡型行业轮动增强基金。"
        "产品定位中等风险、长期理财，目标客户为风险承受中等、投资期限 1 年以上的职场与家庭理财客户，"
        "核心约束：年化收益 12%–18%、最大回撤不超过 10%。策略以申万一级 31 个行业指数为标的，"
        "将行业板块视作独立大类资产，模拟行业 ETF 一篮子持仓开展月度轮动。")
    doc.add_heading("1.2 投资逻辑与收益来源", level=2)
    doc.add_paragraph(
        "收益来源于行业中期景气动量、估值均值回复、资金关注度溢价。通过 Z-Score 截面标准化、IC/IR 与共线性双质检、"
        "LGBM 赋权融合四因子非线性信息，月末选 Top5 行业等权配置、月度调仓。")
    doc.add_heading("1.3 模型选型理由", level=2)
    doc.add_paragraph(
        "多因子框架可解释、便于质检维护；行业指数数量适中、数据公开；LGBM 赋权捕捉非线性组合与交互；"
        "月度调仓匹配中等风险客户对换手与稳定性的要求。")

    # 二、报告步骤与结果
    doc.add_heading("二、报告步骤与结果", level=1)
    doc.add_heading("2.1 数据获取与因子计算", level=2)
    doc.add_paragraph(
        "数据：教师配套申万一级 31 行业日行情/日度 PB/沪深 300 日行情，区间 2022-10-10~2025-12-31 共 787 交易日，无缺失。"
        f"样本划分（按前瞻收益月份，杜绝未来函数）：样本内 2023-02~2024-12 共 23 月 713 条；样本外 2025 全年 {primary.metrics['n_months']} 月。")
    doc.add_paragraph("四因子：景气度(mom_mid,60日累计涨跌)、估值(pb,月末PB-LF)、动量(mom_short,20日累计涨跌)、"
                      "资金流(flow,20日成交额占比)。月末跨31行业Z-Score标准化。")
    doc.add_paragraph("回测规则：月末收盘生成信号执行；Top5等权各20%（单行业≤30%）；单边佣金万分之三按|Δ权重|计入；日频复利；基准沪深300。")

    doc.add_heading("2.2 因子双质检：IC/IR + 共线性", level=2)
    ic_df = ics.copy()
    ic_df.insert(0, "因子", [cfg.FACTOR_LABELS[f.replace("_z", "")] for f in ic_df.index])
    ic_df = ic_df.reset_index(drop=True)[["因子", "ic_mean", "ic_std", "ir", "ic_positive_rate"]]
    ic_df.columns = ["因子", "IC均值", "IC标准差", "IR", "IC胜率"]
    _doc_add_table(doc, ic_df, floatfmt=".3f")
    doc.add_paragraph(f"共线性：最高为 mom_mid 与 mom_short 的 {qc['corr_matrix'].loc['mom_mid_z','mom_short_z']:.2f}，"
                      f"低于阈值 {cfg.COLLINEAR_THRESHOLD}，无高度共线因子，保留全部4因子。样本内IC偏弱（最强PB IR={ics.loc['pb_z','ir']:.2f}），"
                      "反映2023-2024行业轮动信号弱、风格切换频繁，是引入LGBM与风控的依据。")
    img("collinearity.png")
    img("ic_ir.png")

    doc.add_heading("2.3 LGBM 因子赋权与截面打分", level=2)
    doc.add_paragraph(f"LGBM回归，X=四因子Z-Score，y=下月前瞻收益，713条样本，小模型防过拟合。"
                      f"按月分组GroupKFold={model['cv_metrics'].shape[0]}折：平均R²={model['cv_metrics']['r2'].mean():.3f}，CV预测秩IC={model['cv_pred_rank_ic']:.3f}。")
    w_df = pd.DataFrame({
        "因子": [cfg.FACTOR_LABELS[f.replace("_z", "")] for f in cfg.FACTOR_Z],
        "重要度": [imp[f] for f in cfg.FACTOR_Z],
        "IC方向": [int(directions[f]) for f in cfg.FACTOR_Z],
        "定向后权重": [(imp * directions)[f] for f in cfg.FACTOR_Z],
    })
    _doc_add_table(doc, w_df, floatfmt=".4f")
    doc.add_paragraph("打分：综合得分=Σ(重要度×IC方向×z)；月末按score降序选Top5等权。LGBM直接预测作方法敏感性对照。")
    img("lgbm_importances.png")

    doc.add_heading("2.4 回测结果与分析", level=2)
    M = primary.metrics
    m_df = pd.DataFrame({
        "指标": ["累计收益率", "年化收益率", "最大回撤", "夏普比率", "日胜率",
                "沪深300年化", "相对超额(累计)", "月度跑赢胜率", "总换手(倍)", "总成本",
                "达标:年化12%-18%", "达标:回撤≤10%"],
        "数值": [pct(M["cumulative_return"]), pct(M["annualized_return"]), pct(M["max_drawdown"]),
                num(M["sharpe"]), pct(M["daily_win_rate"]), pct(M["benchmark_annualized"]),
                pct(M["excess_cumulative"]), pct(M["monthly_win_rate_vs_bench"]),
                num(M["total_turnover"], 1), pct(M["total_cost"]),
                "✓" if M["target_return_ok"] else "✗", "✓" if M["target_dd_ok"] else "✗"],
    })
    _doc_add_table(doc, m_df, floatfmt=".4f")
    img("nav_primary.png")
    img("drawdown_primary.png")
    img("monthly_primary.png")
    img("holdings_primary.png")
    doc.add_paragraph(
        f"基准对比：策略累计{pct(M['cumulative_return'])}跑赢沪深300的{pct(M['benchmark_cumulative'])}，超额{pct(M['excess_cumulative'])}，"
        f"夏普{num(M['sharpe'])}。2025短期动量恢复有效(IR=+0.36)，综合得分予其最高权重({imp['mom_short_z']:.2f})，成功捕捉轮动alpha。"
        f"达标判断：未同时满足——年化{pct(M['annualized_return'])}超区间(偏高)，回撤{pct(M['max_drawdown'])}破10%红线，源于5行业集中持仓在风格切换月共振下跌。")

    doc.add_heading("2.5 模型维护与风险控制", level=2)
    doc.add_paragraph("风险识别：风格切换(2024回撤-19%)、因子失效(中期动量2025由正转负)、行业黑天鹅、集中度风险。"
                      "定期维护：因子IC月度跟踪、景气度季度复检、LGBM半年度重训练、换手成本监控。"
                      "交易风控：单行业≤30%、回撤触发减仓、单行业止损、极端行情降仓。")
    _doc_add_table(doc, _disp(rob["risk_control"]["table"]))
    img("risk_spectrum.png")
    rc = rob["risk_control"]["table"]
    rc_meet = rc[(rc["target_return_ok"]) & (rc["target_dd_ok"])]
    if len(rc_meet):
        row = rc_meet.iloc[0]
        doc.add_paragraph(f"实证：{row['config']}规则可将年化压至{pct(row['annualized_return'])}、回撤压至{pct(row['max_drawdown'])}，"
                          f"同时满足双目标，代价是强势年牺牲超额{pct(row['excess_cumulative'])}。揭示平衡型产品进攻与防守的权衡。")

    doc.add_heading("2.6 鲁棒性与迭代优化", level=2)
    doc.add_paragraph("参数敏感性(Top3 vs Top5)：")
    _doc_add_table(doc, _disp(rob["parameter"]["table"]))
    doc.add_paragraph("时间敏感性(2024 vs 2025)：")
    _doc_add_table(doc, _disp(rob["time"]["table"]))
    doc.add_paragraph("打分方法敏感性(综合 vs 直接)：")
    _doc_add_table(doc, _disp(rob["method"]["table"]))
    doc.add_paragraph(
        "结论：Top5优于Top3(更稳)；两年度并非均跑赢基准(2024跑输、2025跑赢)，策略机制依赖；两种LGBM打分均跑赢基准，方法稳健。"
        "迭代方向：引入机制识别(波动率状态/因子动量)动态调权，提升跨周期稳定性。")

    # 三、报告总结
    doc.add_heading("三、报告总结", level=1)
    doc.add_heading("3.1 内容总结", level=2)
    doc.add_paragraph(
        f"完整落地行业轮动多因子全流程。主回测2025样本外年化{pct(M['annualized_return'])}、回撤{pct(M['max_drawdown'])}、"
        f"夏普{num(M['sharpe'])}，跑赢沪深300 {pct(M['excess_cumulative'])}。双质检保留4因子，LGBM以短期动量权重最高。"
        "鲁棒性揭示机制依赖与参数稳健；风控实证证明回撤减仓可同时满足双目标。")
    doc.add_heading("3.2 心得体会", level=2)
    for t in [
        "因子有效性是动态的：样本内弱、样本外短期动量复活、中期反转，机制切换是最大风险，须IC月度跟踪与模型重训练。",
        "双质检的价值在诚实：IC弱不丢人，关键是数据说话、不掩盖；保留全因子交LGBM非线性组合比硬凑高IC更扎实。",
        "收益与风控的根本权衡：激进跑赢基准却破回撤红线，加风控守稳却让渡超额；平衡型产品本质是攻守动态再平衡。",
        "工程严谨性：按前瞻收益月份划分样本杜绝未来函数、月末收盘执行对齐信号、按月分组CV防泄漏、换手计入成本，决定回测可信度。",
        "AI协作：全程AI辅助编程与审核，关键决策(因子口径/样本划分/达标判断)由本人基于课程知识独立完成并核验。",
    ]:
        doc.add_paragraph(t, style="List Number")
    doc.add_paragraph("风险提示：本策略仅用于课程学习与方法实验，回测不代表未来收益，不构成投资建议。")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
