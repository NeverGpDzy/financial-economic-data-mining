"""
Fill the homework 4B AI code audit Word template with the reviewed results.
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT.parent / "作业四-B" / "AI代码审查与修复表 学生表.docx"
OUTPUT_DIR = ROOT / "outputs" / "homework4b"
OUTPUT_PATH = OUTPUT_DIR / "202331060205_丁致宇_AI代码审查与修复表.docx"


def fill_cell(cell, text, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)


def fill_review_row(table, row, location, issue, severity, status, fix):
    fill_cell(table.cell(row, 2), location)
    fill_cell(table.cell(row, 3), issue)
    fill_cell(table.cell(row, 4), severity)
    fill_cell(table.cell(row, 5), status)
    fill_cell(table.cell(row, 6), fix)


def main():
    doc = Document(str(TEMPLATE_PATH))

    info = doc.tables[0]
    fill_cell(info.cell(1, 1), "丁致宇")
    fill_cell(info.cell(2, 1), "202331060205")
    fill_cell(info.cell(3, 1), "作业4B：中频短线量化全流程（AI辅助版）")
    fill_cell(info.cell(4, 1), "☑ 日频短线  □ 周频中线  □ 混合周期")
    fill_cell(info.cell(5, 1), "□ 纯AI生成  ☑ AI生成+人工修改  □ 人工自主编写")
    fill_cell(info.cell(6, 1), "约70%")
    fill_cell(info.cell(7, 1), "□ 纯快变量  ☑ 快变量为主+慢变量过滤  □ 快慢变量混合")

    fill_review_row(
        doc.tables[1], 1,
        "data.py:compute_factors",
        "因子使用T日收盘后已知数据，预测T+1日收益；未使用未来收益计算当日信号。",
        "无",
        "合规",
        "Reversal=-ret，标签next_excess_ret仅用于训练/质检，不参与当日信号生成。",
    )
    fill_review_row(
        doc.tables[1], 2,
        "data.py:load_parquet_data",
        "数据加载后按ts_code和trade_date排序，信号计算与回测按时间顺序执行。",
        "无",
        "合规",
        "df = df.sort_values(['ts_code', 'trade_date']).reset_index(drop=True)",
    )
    fill_review_row(
        doc.tables[1], 3,
        "backtest.py:run_backtest",
        "需确认信号与收益结算时间错开，避免当日信号当日收益。",
        "无",
        "合规",
        "T日收盘后按LGBM得分选股调仓，T+1用持仓股票ret结算收益。",
    )
    fill_review_row(
        doc.tables[1], 4,
        "main.py:run_factor_quality_check",
        "原AI初稿对full_df做因子质检，且3σ缩尾曾使用全样本阈值，存在样本外信息污染风险。",
        "P0",
        "已修复",
        "仅对2020-2023训练集计算IC/IR、VIF和OLS；3σ缩尾改为交易日截面处理，2024-2025只用于回测。",
    )
    fill_review_row(
        doc.tables[1], 5,
        "models.py:train_lgbm_with_cv",
        "原时间序列CV名义5折，但第一折无训练集被跳过，实际不足5个验证窗口。",
        "P1",
        "已修复",
        "改为6个时间块生成5个扩展窗口CV，最佳Fold 5，验证MSE=0.000410。",
    )

    fill_review_row(
        doc.tables[2], 1,
        "config.py:FACTOR_NAMES",
        "原AI初稿额外加入Momentum、Volatility、Turnover、VolumeChange，偏离题目4因子要求。",
        "P1",
        "已修复",
        "最终只保留Reversal、Liquidity、MoneyFlow、Value四个标准因子。",
    )
    fill_review_row(
        doc.tables[2], 2,
        "data.py:compute_factors",
        "Value为PE_TTM倒数，属于慢变量；在本作业中仅作为四因子之一参与截面排序辅助。",
        "无",
        "合规",
        "Value=1/pe_ttm，低估值正向化；未使用未来财务数据。",
    )
    fill_review_row(
        doc.tables[2], 3,
        "data.py:clean_data",
        "原清洗阶段丢弃next_excess_ret为空的最后交易日，导致回测无法覆盖2025-12-31。",
        "P1",
        "已修复",
        "清洗只剔除ret为空记录；标签缺失只在建模和IC环节过滤。",
    )
    fill_review_row(
        doc.tables[2], 4,
        "factors.py:compute_vif",
        "需验证四因子是否存在高共线。",
        "无",
        "合规",
        "VIF最大为MoneyFlow的1.1039，远低于VIF>5阈值。",
    )
    fill_review_row(
        doc.tables[2], 5,
        "data.py:compute_factors",
        "本地数据缺少个股级buy_amount/sell_amount，MoneyFlow无法按原字段逐字实现。",
        "P1",
        "已修复",
        "使用5日滚动签名成交额净流入/circ_mv作为个股级代理；反转、流动性、资金流、价值均有经济含义。",
    )

    fill_review_row(
        doc.tables[3], 1,
        "config.py:COMMISSION_RATE",
        "交易成本需纳入回测，避免零成本理想化。",
        "P1",
        "已修复",
        "单边成本0.1%；调仓按卖出+买入实际换手部分扣成本。",
    )
    fill_review_row(
        doc.tables[3], 2,
        "config.py:TOP_N_STOCKS",
        "策略为Top5等权，单票20%。总仓位100%符合题目规则，但实盘风险敞口偏高。",
        "P2",
        "保留",
        "作业要求Top5等权；后续实盘可增加总仓位上限或波动率目标仓位。",
    )
    fill_review_row(
        doc.tables[3], 3,
        "backtest.py",
        "当前无显式止损机制，主要依靠每日调仓和模型重新打分控制风险。",
        "P2",
        "待优化",
        "本次按作业要求完成回测；后续可加入单票止损和账户回撤控制。",
    )
    fill_review_row(
        doc.tables[3], 4,
        "data.py:load_parquet_data",
        "标的池需满足流动性，避免无法成交。",
        "无",
        "合规",
        "使用沪深300成分股，属于A股流动性较优股票池。",
    )
    fill_review_row(
        doc.tables[3], 5,
        "config.py:HOLDING_DAYS",
        "需符合1~2日中频短线持仓规则。",
        "无",
        "合规",
        "HOLDING_DAYS=1，每日调仓，平均持仓1~2日。",
    )

    fill_review_row(
        doc.tables[4], 1,
        "requirements.txt / imports",
        "需检查是否存在AI虚构库或不存在API。",
        "无",
        "合规",
        "已验证pandas、numpy、scipy、lightgbm、python-pptx等依赖可导入。",
    )
    fill_review_row(
        doc.tables[4], 2,
        "backtest.py",
        "需检查买卖逻辑是否出现同时冲突或重复扣费。",
        "无",
        "合规",
        "先结算旧持仓收益，再按新旧持仓差异扣买卖成本，逻辑清晰。",
    )
    fill_review_row(
        doc.tables[4], 3,
        "models.py",
        "机器学习模型需避免过拟合和未来数据泄露。",
        "无",
        "合规",
        "使用时间序列扩展窗口CV，验证窗口始终位于训练窗口之后。",
    )
    fill_review_row(
        doc.tables[4], 4,
        "report/build_homework4b_ppt.py",
        "原PPT脚本写死旧收益和8因子结论，重跑后会展示错误数字。",
        "P1",
        "已修复",
        "PPT脚本改为读取最新CSV/JSON自动生成，旧版硬编码收益数字已移除。",
    )
    fill_review_row(
        doc.tables[4], 5,
        "config.py / data.py",
        "原AI加入额外因子以提升表现，属于与题目不一致的冗余逻辑。",
        "P1",
        "已修复",
        "删除扩展因子计算和展示，输出数据表也无扩展因子列。",
    )

    fill_review_row(
        doc.tables[5], 1,
        "全局",
        "核心模块有函数划分和中文说明。",
        "无",
        "合规",
        "保留模块化结构：data、factors、models、backtest、plots、main。",
    )
    fill_review_row(
        doc.tables[5], 2,
        "全局",
        "变量和函数命名整体符合Python规范。",
        "无",
        "合规",
        "使用snake_case命名，关键参数集中在config.py。",
    )
    fill_review_row(
        doc.tables[5], 3,
        "全局",
        "流程结构与作业模块一致。",
        "无",
        "合规",
        "本地数据加载→因子计算→质检→LGBM→样本外回测→图表/PPT。",
    )
    fill_review_row(
        doc.tables[5], 4,
        "data.py / factors.py",
        "部分函数仍以数据清洗和NaN过滤为主，异常捕获不算完整。",
        "P2",
        "部分修复",
        "当前已能稳定重跑；后续可增加更细的文件缺失和字段缺失错误提示。",
    )
    fill_review_row(
        doc.tables[5], 5,
        "config.py:LGBM_PARAMS",
        "结果需要可复现。",
        "无",
        "合规",
        "LightGBM设置seed=42，重跑结果可复现。",
    )

    overall = doc.tables[6]
    fill_cell(
        overall.cell(1, 1),
        "1. 因子口径偏离作业要求：原AI加入4个扩展因子，已改回4因子。\n"
        "2. 样本外信息污染：原IC/OLS使用full_df、缩尾使用全样本阈值，已改为训练集质检和按日截面缩尾。\n"
        "3. 回测细节偏差：首日成本、末日调仓、基准起点已修复。",
    )
    fill_cell(
        overall.cell(2, 1),
        "1. 统一4因子口径并重跑全部数据。\n"
        "2. 修正训练/回测隔离，样本外只用于2024-2025打分与回测。\n"
        "3. 重做PPT，直接读取最新结果，避免旧版硬编码。",
    )
    fill_cell(
        overall.cell(3, 1),
        "□ 完全合规（所有P0/P1问题已修复）\n"
        "☑ 基本合规（P0/P1已修复，仅剩止损和仓位上限等P2实盘优化项）\n"
        "□ 不合规（存在未修复的致命缺陷）",
    )
    fill_cell(
        overall.cell(4, 1),
        "AI适合快速生成量化流程框架，但会倾向于添加额外因子或复用全样本结果来提高表现。"
        "本次人工审查重点修复了题目口径、样本隔离、交易成本和展示材料一致性。"
        "量化作业不能只看收益高低，必须确认每个数字都来自正确数据区间和正确交易假设。",
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"AI代码审查与修复表已保存至: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
